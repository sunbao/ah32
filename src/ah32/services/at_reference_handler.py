"""@引用处理器 - 阿蛤（AH32）核心功能

根据 docs/AH32_DESIGN.md 设计：
- 支持@路径引用语法
- 自动读取引用文件并向量入库
- 供RAG检索使用

示例：
"参考@D:\\资料\\合同模板.docx检查这个合同"
"根据@D:\\知识库\\技术规范.md补充第三章"
"""

from __future__ import annotations

import re
import os
import logging
import hashlib
import time
import tempfile
from typing import List, Dict, Any, Optional
from pathlib import Path

from ah32.knowledge.store import LocalVectorStore

# 检查win32com是否可用
try:
    import win32com.client as win32
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False

logger = logging.getLogger(__name__)


class AtReferenceHandler:
    """@引用处理器"""

    def __init__(self, vector_store: LocalVectorStore):
        self.vector_store = vector_store

    def extract_at_references(self, user_input: str) -> List[str]:
        """提取@引用的文件路径

        Args:
            user_input: 用户输入文本

        Returns:
            提取的文件路径列表

        Examples:
            >>> handler = AtReferenceHandler(vector_store)
            >>> paths = handler.extract_at_references("参考@D:\\资料\\合同.docx检查")
            >>> print(paths)
            ['D:\\资料\\合同.docx']
        """
        logger.debug(f"开始提取@引用，输入长度: {len(user_input)}")
        logger.debug(f"用户输入内容: {user_input[:200]}...")

        # 正则表达式：匹配@后面跟路径
        pattern = r'@([^\s\[\]{}()（）]+)'
        matches = re.findall(pattern, user_input)
        logger.debug(f"正则匹配结果: {len(matches)} 个匹配项")

        # 清理路径
        paths = []
        for match in matches:
            # 移除可能的末尾标点
            cleaned = re.sub(r'[，。；：！]$', '', match)
            logger.debug(f"路径清理: '{match}' -> '{cleaned}'")
            paths.append(cleaned)

        logger.info(f"提取到{len(paths)}个@引用路径: {paths}")
        return paths

    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件内容哈希值

        Args:
            file_path: 文件路径

        Returns:
            文件内容的MD5哈希值
        """
        try:
            with open(file_path, 'rb') as f:
                return hashlib.md5(f.read()).hexdigest()
        except Exception as e:
            logger.error(f"计算文件哈希失败 {file_path}: {e}")
            return ""

    async def handle(self, user_input: str) -> Dict[str, Any]:
        """处理用户输入中的@引用

        Args:
            user_input: 用户输入文本

        Returns:
            {
                "paths": [...],  # 提取的路径列表
                "processed": [...],  # 已入库的文件
                "errors": [...]  # 处理错误
            }
        """
        logger.debug(f"开始处理@引用，输入: {user_input[:100]}...")
        try:
            # 1. 提取@引用路径
            logger.debug("步骤1: 提取@引用路径")
            paths = self.extract_at_references(user_input)
            logger.debug(f"提取结果: {len(paths)} 个路径")

            if not paths:
                logger.debug("未找到@引用路径，返回空结果")
                return {
                    "paths": [],
                    "processed": [],
                    "errors": []
                }

            # 2. 处理每个文件
            logger.debug(f"步骤2: 开始处理 {len(paths)} 个文件")
            processed = []
            errors = []

            for path_idx, path in enumerate(paths):
                logger.debug(f"处理文件 {path_idx + 1}/{len(paths)}: {path}")
                try:
                    # 检查文件是否存在
                    if not os.path.exists(path):
                        error_msg = f"文件不存在: {path}"
                        logger.warning(error_msg)
                        errors.append(error_msg)
                        continue

                    file_size = os.path.getsize(path)
                    logger.debug(f"文件存在，大小: {file_size} bytes")

                    # 读取文件内容
                    logger.debug("步骤2.1: 读取文件内容")
                    content = await self._read_file(path)
                    if not content:
                        error_msg = f"无法读取文件: {path}"
                        logger.warning(error_msg)
                        errors.append(error_msg)
                        continue

                    logger.debug(f"文件读取成功，内容长度: {len(content)} 字符")

                    # 智能入库（去重 + 更新检测）
                    logger.debug("步骤2.2: 智能入库（去重+更新检测）")
                    result = await self._upsert_document(path, content)
                    logger.debug(f"入库结果: status={result['status']}, action={result['action']}, updated={result['updated']}")

                    processed.append({
                        "path": path,
                        "content_length": len(content),
                        "status": result["status"],
                        "action": result["action"],
                        "updated": result["updated"]
                    })

                    if result["updated"]:
                        logger.info(f"成功{result['action']}@引用文件: {path} (新增{result['chunks_added']}个块)")
                    else:
                        logger.info(f"文件未变化，跳过: {path}")

                except Exception as e:
                    error_msg = f"处理文件失败 {path}: {str(e)}"
                    logger.error(error_msg)
                    logger.debug(f"错误详情: {type(e).__name__}: {str(e)}", exc_info=True)
                    errors.append(error_msg)

            logger.debug(f"@引用处理完成: 处理{len(paths)}个文件，成功{len(processed)}个，失败{len(errors)}个")
            return {
                "paths": paths,
                "processed": processed,
                "errors": errors
            }

        except Exception as e:
            logger.error(f"@引用处理失败: {str(e)}")
            logger.debug(f"错误详情: {type(e).__name__}: {str(e)}", exc_info=True)
            return {
                "paths": [],
                "processed": [],
                "errors": [str(e)]
            }

    async def _read_file(self, file_path: str) -> Optional[str]:
        """读取文件内容

        Args:
            file_path: 文件路径

        Returns:
            文件内容文本
        """
        try:
            file_path = Path(file_path)
            suffix = file_path.suffix.lower()

            if suffix in ['.docx', '.doc']:
                # 读取Word文档
                try:
                    from docx import Document

                    # 处理.docx和.doc格式
                    if suffix == '.docx':
                        # .docx文件直接读取
                        logger.debug(f"使用python-docx读取.docx文档: {file_path}")
                        doc = Document(file_path)
                        temp_file = None

                    elif suffix == '.doc':
                        # .doc文件需要先转换
                        logger.debug(f"检测到.doc文件，开始转换: {file_path}")

                        if not HAS_WIN32COM:
                            raise ValueError(
                                "处理.doc格式需要安装pywin32库。\n"
                                "请运行: pip install pywin32"
                            )

                        # 转换为.docx格式
                        temp_file = self._convert_doc_to_docx(file_path)
                        logger.debug(f"转换完成，使用临时文件: {temp_file}")
                        doc = Document(temp_file)

                    # 提取文本内容
                    content_parts = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            content_parts.append(paragraph.text.strip())

                    # 提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text.strip())
                            if row_text:
                                content_parts.append(" | ".join(row_text))

                    content = "\n".join(content_parts)

                    # 清理临时文件
                    if temp_file and temp_file.exists():
                        try:
                            temp_file.unlink()
                            logger.debug(f"临时文件已清理: {temp_file}")
                        except Exception as e:
                            logger.warning(f"清理临时文件失败 {temp_file}: {e}")

                    logger.debug(f"Word文档读取成功，内容长度: {len(content)} 字符")
                    return content

                except ImportError:
                    logger.warning("python-docx库未安装，无法读取Word文档")
                    return None
                except Exception as e:
                    logger.error(f"读取Word文档失败 {file_path}: {e}")
                    # 确保清理临时文件
                    if 'temp_file' in locals() and temp_file and temp_file.exists():
                        try:
                            temp_file.unlink()
                        except Exception as e2:
                            logger.warning(f"清理临时文件失败 {temp_file}: {e2}", exc_info=True)
                    raise ValueError(f"读取{file_path.suffix}文件失败: {str(e)}")

            elif suffix in ['.md', '.txt']:
                # 读取文本文件
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            else:
                logger.info(f"不支持的文件类型: {suffix}")
                return None

        except Exception as e:
            logger.error(f"读取文件失败 {file_path}: {str(e)}")
            return None

    def _convert_doc_to_docx(self, doc_path: Path) -> Path:
        """将.doc格式转换为.docx格式

        Args:
            doc_path: .doc文件路径

        Returns:
            转换后的.docx文件路径（临时文件）

        Raises:
            ValueError: 当win32com不可用或转换失败时
        """
        if not HAS_WIN32COM:
            raise ValueError("需要安装pywin32库才能处理.doc格式文件")

        try:
            import win32com.client as win32
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            # 打开.doc文件
            doc = word.Documents.Open(str(doc_path.absolute()))

            # 创建临时.docx文件
            temp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
            temp_docx.close()

            # 保存为.docx格式 (wdFormatXMLDocument = 12)
            doc.SaveAs(temp_docx.name, FileFormat=12)
            doc.Close()
            word.Quit()

            logger.info(f"已将.doc文件转换为临时.docx: {temp_docx.name}")
            return Path(temp_docx.name)

        except Exception as e:
            error_msg = f"转换.doc文件失败: {str(e)}"
            logger.error(error_msg)
            raise ValueError(error_msg)

    async def _upsert_document(self, file_path: str, content: str) -> Dict[str, Any]:
        """智能文档更新（去重 + 更新检测）

        Args:
            file_path: 文件路径
            content: 文件内容

        Returns:
            {
                "updated": bool,  # 是否进行了更新
                "action": str,    # 操作类型：'added', 'updated', 'skipped'
                "status": str,    # 状态：'success', 'skipped', 'error'
                "chunks_added": int  # 新增的块数量
            }
        """
        logger.debug(f"开始智能入库: {file_path}, 内容长度: {len(content)}")
        try:
            # 1. 计算当前文件哈希
            logger.debug("步骤1: 计算文件哈希")
            current_hash = self._calculate_file_hash(file_path)
            logger.debug(f"文件哈希值: {current_hash[:16]}...")
            if not current_hash:
                logger.error("文件哈希计算失败")
                return {
                    "updated": False,
                    "action": "failed",
                    "status": "error",
                    "chunks_added": 0
                }

            # 2. 检查文件是否已存在（使用优化后的方法）
            logger.debug("步骤2: 检查文件是否存在")
            file_exists = self.vector_store.check_document_exists(file_path)
            logger.debug(f"文件存在状态: {file_exists}")

            if file_exists:
                # 获取已存在文档的哈希值
                logger.debug("获取已存在文档的哈希值")
                existing_docs = self.vector_store.get_documents_by_source(file_path, limit=1)
                if existing_docs:
                    stored_hash = existing_docs[0].get("metadata", {}).get("file_hash")
                    logger.debug(f"已存储哈希: {stored_hash[:16] if stored_hash else 'None'}...")

                    if stored_hash == current_hash:
                        # 文件未变化，跳过
                        logger.info(f"文件未变化，跳过: {file_path}")
                        return {
                            "updated": False,
                            "action": "skipped",
                            "status": "success",
                            "chunks_added": 0
                        }
                    else:
                        # 文件已更新，删除旧版本
                        logger.info(f"检测到文件更新，删除旧版本: {file_path}")
                        logger.debug(f"旧哈希: {stored_hash[:16]}..., 新哈希: {current_hash[:16]}...")
                        self.vector_store.delete_by_filter({"source": file_path})
                        logger.debug("旧版本文档已删除")
                else:
                    logger.debug("无法获取已存在文档信息，将作为新文档处理")

            # 3. 分块处理
            logger.debug("步骤3: 分块处理")
            chunks = self._split_into_chunks(content, max_length=1000)
            logger.debug(f"分块完成，共 {len(chunks)} 个块")

            # 4. 为每个块创建文档对象（带哈希值）
            logger.debug("步骤4: 创建文档对象")
            from langchain_core.documents import Document

            documents = []
            for i, chunk in enumerate(chunks):
                logger.debug(f"创建块 {i+1}/{len(chunks)}, 长度: {len(chunk)}")
                doc = Document(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_index": i,
                        "file_name": os.path.basename(file_path),
                        "type": "at_reference",
                        "file_hash": current_hash,  # 添加文件哈希
                        "updated_at": time.time(),
                        "first_added": time.time() if not file_exists else None
                    }
                )
                documents.append(doc)

            # 5. 添加到向量存储
            logger.debug("步骤5: 添加到向量存储")
            if documents:
                logger.debug(f"准备添加 {len(documents)} 个文档块到向量存储")
                self.vector_store.add_documents(documents)
                logger.debug("向量存储添加完成")
                action = "added" if not file_exists else "updated"
                logger.info(f"已{action} {len(documents)}个文档块: {file_path}")
                return {
                    "updated": True,
                    "action": action,
                    "status": "success",
                    "chunks_added": len(documents)
                }
            else:
                logger.warning("没有文档需要添加")

            return {
                "updated": False,
                "action": "failed",
                "status": "error",
                "chunks_added": 0
            }

        except Exception as e:
            logger.error(f"智能入库失败 {file_path}: {str(e)}")
            logger.debug(f"入库错误详情: {type(e).__name__}: {str(e)}", exc_info=True)
            return {
                "updated": False,
                "action": "error",
                "status": "error",
                "chunks_added": 0
            }

    # 保留原方法以保持兼容性（已弃用）
    async def _add_to_vector_store(self, file_path: str, content: str):
        """添加到向量存储（已弃用，使用_upsert_document替代）

        Args:
            file_path: 文件路径
            content: 文件内容
        """
        logger.warning("_add_to_vector_store 已弃用，请使用 _upsert_document")
        # 调用新方法
        result = await self._upsert_document(file_path, content)
        if not result["updated"]:
            raise Exception(f"入库失败: {result['action']}")

    def _split_into_chunks(self, content: str, max_length: int = 1000) -> List[str]:
        """将长文本分块

        Args:
            content: 文本内容
            max_length: 最大块长度

        Returns:
            文本块列表
        """
        if len(content) <= max_length:
            return [content]

        chunks = []
        current_chunk = ""

        # 按段落分割
        paragraphs = content.split('\n\n')

        for paragraph in paragraphs:
            if len(current_chunk) + len(paragraph) + 2 <= max_length:
                if current_chunk:
                    current_chunk += '\n\n' + paragraph
                else:
                    current_chunk = paragraph
            else:
                # 保存当前块
                if current_chunk:
                    chunks.append(current_chunk)

                # 如果段落本身太长，需要进一步分割
                if len(paragraph) > max_length:
                    # 按句号分割
                    sentences = paragraph.split('。')
                    temp_chunk = ""

                    for sentence in sentences:
                        if len(temp_chunk) + len(sentence) + 1 <= max_length:
                            if temp_chunk:
                                temp_chunk += '。' + sentence
                            else:
                                temp_chunk = sentence
                        else:
                            if temp_chunk:
                                chunks.append(temp_chunk)
                            temp_chunk = sentence

                    if temp_chunk:
                        current_chunk = temp_chunk
                    else:
                        current_chunk = ""
                else:
                    current_chunk = paragraph

        # 添加最后一个块
        if current_chunk:
            chunks.append(current_chunk)

        return chunks


# 便捷函数
async def handle_at_references(user_input: str, vector_store: LocalVectorStore) -> Dict[str, Any]:
    """处理@引用的便捷函数"""
    handler = AtReferenceHandler(vector_store)
    return await handler.handle(user_input)


def extract_at_paths(user_input: str) -> List[str]:
    """提取@路径的便捷函数"""
    handler = AtReferenceHandler(None)  # 不需要vector_store
    return handler.extract_at_references(user_input)
