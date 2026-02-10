"""RAG知识库API - 文档导入、查询、删除等操作"""

from __future__ import annotations

from typing import List, Dict, Any, Optional, Set, Tuple
from pathlib import Path
import os
import json
import logging
from datetime import datetime
import hashlib
import time
import tempfile
import asyncio
import uuid

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

logger = logging.getLogger(__name__)

# Project-scoped RAG (project + global) without duplicating vectors.
from ah32.knowledge.project_rag_index import (
    ProjectRagIndex,
    ProjectContext,
    derive_project_context,
    GLOBAL_PROJECT_ID,
)
from ah32.config import settings

# 创建路由
router = APIRouter(prefix="/api/rag", tags=["RAG知识库"])

# ----------------------------
# Shared path mode: client path -> backend path mapping
# ----------------------------

def _parse_path_mappings(raw: str) -> List[Tuple[str, str]]:
    """Parse mappings like:
    - Z:\\share=>/mnt/share
    - \\\\server\\share\\=>/mnt/share/
    Multiple entries separated by newlines or ';'.
    """
    if not raw:
        return []
    parts: List[str] = []
    for line in raw.replace("\r", "\n").split("\n"):
        line = line.strip()
        if line:
            parts.append(line)
    for seg in raw.split(";"):
        seg = seg.strip()
        if seg and seg not in parts:
            parts.append(seg)

    out: List[Tuple[str, str]] = []
    for part in parts:
        if "=>" in part:
            a, b = part.split("=>", 1)
        elif "=" in part:
            a, b = part.split("=", 1)
        else:
            continue
        a = a.strip()
        b = b.strip()
        if a and b:
            out.append((a, b))

    # De-dupe while keeping order.
    seen = set()
    uniq: List[Tuple[str, str]] = []
    for a, b in out:
        k = (a, b)
        if k in seen:
            continue
        seen.add(k)
        uniq.append((a, b))
    return uniq


def _load_path_mappings() -> List[Tuple[str, str]]:
    raw = (os.environ.get("AH32_PATH_MAPPINGS") or os.environ.get("PATH_MAPPINGS") or "").strip()
    return _parse_path_mappings(raw)


def _normalize_client_path_for_match(p: str) -> str:
    s = str(p or "").strip()
    if not s:
        return ""
    # Normalize to backslashes for matching; keep UNC semantics.
    s = s.replace("/", "\\")
    return s.lower()


def _ensure_trailing_sep(p: str) -> str:
    s = str(p or "").strip()
    if not s:
        return s
    if s.endswith("\\") or s.endswith("/"):
        return s
    # Windows-like when it looks like drive or UNC.
    if s.startswith("\\\\") or (len(s) >= 2 and s[1] == ":"):
        return s + "\\"
    return s + "/"


def _apply_path_mapping(client_path: str) -> Tuple[str, Optional[Dict[str, str]]]:
    """Return (backend_path, mapping_info_or_none)."""
    cp = str(client_path or "").strip()
    if not cp:
        return cp, None

    mappings = _load_path_mappings()
    if not mappings:
        return cp, None

    cp_norm = _normalize_client_path_for_match(cp)
    for raw_from, raw_to in mappings:
        from_prefix = _ensure_trailing_sep(raw_from)
        to_prefix = _ensure_trailing_sep(raw_to)
        from_norm = _normalize_client_path_for_match(from_prefix)
        if from_norm and cp_norm.startswith(from_norm):
            rest = cp[len(from_prefix) :]
            if to_prefix.startswith("/"):
                backend = to_prefix.rstrip("/") + "/" + rest.replace("\\", "/").lstrip("/")
            else:
                backend = to_prefix.rstrip("\\") + "\\" + rest.replace("/", "\\").lstrip("\\")
            return backend, {"from": raw_from, "to": raw_to, "mapped": backend}

    return cp, None

# 全局向量存储引用（从main.py获取）
def get_vector_store():
    """获取全局向量存储实例"""
    from ah32.server.main import app
    if hasattr(app.state, '_vector_store'):
        return app.state._vector_store
    raise HTTPException(status_code=500, detail="向量存储未初始化")


def _get_project_index(vector_store) -> ProjectRagIndex:
    """Get (and cache) the project/global membership index next to the vector store."""
    from ah32.server.main import app

    idx = getattr(app.state, "_project_rag_index", None)
    try:
        persist_dir = Path(getattr(vector_store, "persist_path"))
    except Exception:
        persist_dir = None

    needs_reload = True
    try:
        needs_reload = not idx or not persist_dir or Path(getattr(idx, "persist_dir", "")) != persist_dir
    except Exception:
        needs_reload = True

    if needs_reload:
        if not persist_dir:
            # Fallback: store index under temp-less current working directory.
            persist_dir = Path.cwd() / ".ah32_rag"
        idx = ProjectRagIndex.load(persist_dir)

        # Backward compatible bootstrap: if index is missing OR empty but vectors already exist,
        # mark existing sources as global so listing/statistics work out-of-the-box.
        try:
            metas: List[Dict[str, Any]] = []
            try:
                metas = vector_store.get_all_metadatas()
            except Exception:
                metas = []
            sources = sorted({(m or {}).get("source") for m in (metas or []) if (m or {}).get("source")})
            if sources and (not idx.index_path.exists() or idx.is_empty()):
                idx.bootstrap_all_as_global(sources)
                idx.save()
        except Exception as e:
            logger.warning(f"[rag] project index bootstrap failed: {e}", exc_info=True)

        app.state._project_rag_index = idx

    return idx


def _resolve_project_scope(
    *,
    vector_store,
    project_id: Optional[str],
    context_document_path: Optional[str],
) -> Tuple[Optional[str], Optional[str]]:
    """Resolve (project_id, project_label) from explicit id or active document path."""
    pid = (project_id or "").strip()
    if pid:
        if pid == GLOBAL_PROJECT_ID:
            return GLOBAL_PROJECT_ID, "全局库"
        idx = _get_project_index(vector_store)
        label = idx.get_project_label(pid) or pid
        return pid, label

    ctx_path = (context_document_path or "").strip()
    if not ctx_path:
        return None, None

    ctx = derive_project_context(ctx_path)
    if not ctx:
        return None, None
    return ctx.project_id, ctx.label

# Pydantic模型
class RagStatisticsResponse(BaseModel):
    totalDocuments: int
    totalVectors: int
    storageSize: str
    importMethods: Dict[str, Dict[str, Any]]
    # Optional: scoped view meta (current project + global).
    scopeProjectId: Optional[str] = None
    scopeProjectLabel: Optional[str] = None
    includeGlobal: bool = True
    globalDocuments: Optional[int] = None
    projectDocuments: Optional[int] = None

class RagDocument(BaseModel):
    name: str
    path: str
    # Optional: a more user-friendly path (e.g. original client path in shared-path mode).
    displayPath: Optional[str] = None
    size: str
    importMethod: str
    importTime: str
    vectors: int
    hash: Optional[str] = None
    # Optional: scope hint for UI.
    scope: Optional[str] = None  # global/project/both/unscoped
    projectsCount: Optional[int] = None

class DocumentsByMethodResponse(BaseModel):
    method: str
    name: str
    count: int
    documents: List[RagDocument]

class DeleteDocumentRequest(BaseModel):
    documentPath: str

class ReimportDocumentRequest(BaseModel):
    documentPath: str
    documentName: str

class ApiResponse(BaseModel):
    success: bool
    data: Optional[Any] = None
    message: Optional[str] = None
    error: Optional[str] = None

class RagStatisticsApiResponse(ApiResponse):
    data: Optional[RagStatisticsResponse] = None

class DocumentsByMethodApiResponse(ApiResponse):
    data: Optional[List[DocumentsByMethodResponse]] = None

class BatchImportDocumentItem(BaseModel):
    # Path-based import (local/shared-path mode). For upload mode, keep empty and provide `text`.
    path: str = ""
    name: str
    # Upload-to-RAG mode (remote default): the client provides plain text to embed.
    text: Optional[str] = None
    # Optional: display alias (e.g. original client path) for upload mode.
    pathAlias: Optional[str] = None
    # wps/agent/api/command/atref
    importMethod: str = "wps"

class BatchImportRequest(BaseModel):
    # Backward-compatible: either provide explicit documents list OR provide a directory to scan.
    documents: Optional[List[BatchImportDocumentItem]] = None
    directory: Optional[str] = None
    recursive: bool = True
    limit: int = 2000
    # Default import method used when scanning a directory (wps/agent/api/command/atref)
    importMethod: str = "agent"
    # Optional scope control:
    # - "global": add files into the global library
    # - "project": add files into the current project (derived from each file path)
    scope: Optional[str] = None
    # If provided, force the project id instead of deriving from the file path.
    projectId: Optional[str] = None
    # For listing/statistics endpoints, frontend can pass the active document path to resolve project label.
    contextDocumentPath: Optional[str] = None

class ValidatePathRequest(BaseModel):
    path: str

class ScanDirectoryRequest(BaseModel):
    directory: str
    recursive: bool = True
    limit: int = 200

class UploadTextItem(BaseModel):
    """Text payload uploaded from the client (e.g. WPS export)."""

    name: str
    text: str
    pathAlias: Optional[str] = None
    importMethod: str = "upload"
    scope: str = "global"  # global|project
    contextDocumentPath: Optional[str] = None


class UploadTextRequest(BaseModel):
    items: List[UploadTextItem]

# =======================
# Task Models / In-Memory Task Registry
# =======================

class RagTaskState(BaseModel):
    taskId: str
    type: str  # batch-import / reimport
    status: str  # pending/running/completed/failed/cancelled
    progress: int = 0
    currentStep: str = "等待执行"
    createdAt: float = 0.0
    updatedAt: float = 0.0
    result: Optional[Dict[str, Any]] = None
    error: Optional[str] = None


_rag_tasks: Dict[str, RagTaskState] = {}
_rag_task_queues: Dict[str, asyncio.Queue] = {}
_rag_async_tasks: Dict[str, asyncio.Task] = {}


def _now_ts() -> float:
    return time.time()


def _task_emit(task_id: str, payload: Dict[str, Any]) -> None:
    """Push a progress/event payload to task queue if present."""
    q = _rag_task_queues.get(task_id)
    if not q:
        return
    try:
        q.put_nowait(payload)
    except Exception:
        # Best-effort; don't block ingestion.
        pass


def _task_update(task_id: str, **updates) -> None:
    task = _rag_tasks.get(task_id)
    if not task:
        return
    for k, v in updates.items():
        setattr(task, k, v)
    task.updatedAt = _now_ts()
    _task_emit(task_id, {"type": "task", "data": task.model_dump()})


# 工具函数
def _format_file_size(size_bytes: int) -> str:
    """格式化文件大小"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    return f"{size_bytes:.1f} {size_names[i]}"

def _get_file_stats(file_path: str) -> Dict[str, Any]:
    """获取文件统计信息"""
    try:
        backend_path, _mapping = _apply_path_mapping(file_path)
        path = Path(backend_path)
        if path.exists():
            stat = path.stat()
            return {
                "size": _format_file_size(stat.st_size),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            }
    except Exception as e:
        logger.warning(f"获取文件统计信息失败 {file_path}: {e}")
    return {"size": "0 B", "modified": "未知"}

def _calculate_file_hash(file_path: str) -> str:
    """计算文件哈希，用于去重/更新检测"""
    try:
        md5 = hashlib.md5()
        backend_path, _mapping = _apply_path_mapping(file_path)
        with open(backend_path, "rb") as f:
            while True:
                chunk = f.read(1024 * 1024)  # 1MB chunks; keep memory bounded for large docs.
                if not chunk:
                    break
                md5.update(chunk)
        return md5.hexdigest()
    except Exception as e:
        logger.warning(f"计算文件哈希失败 {file_path}: {e}")
        return ""

def _split_into_chunks(content: str, max_length: int = 1000) -> List[str]:
    """将长文本分块（与@引用处理保持一致）"""
    if not content:
        return []
    if len(content) <= max_length:
        return [content]

    chunks: List[str] = []
    current_chunk = ""

    paragraphs = content.split("\n\n")
    for paragraph in paragraphs:
        if len(current_chunk) + len(paragraph) + 2 <= max_length:
            current_chunk = f"{current_chunk}\n\n{paragraph}" if current_chunk else paragraph
            continue

        if current_chunk:
            chunks.append(current_chunk)
            current_chunk = ""

        if len(paragraph) > max_length:
            sentences = paragraph.split("。")
            temp_chunk = ""
            for sentence in sentences:
                if len(temp_chunk) + len(sentence) + 1 <= max_length:
                    temp_chunk = f"{temp_chunk}。{sentence}" if temp_chunk else sentence
                else:
                    if temp_chunk:
                        chunks.append(temp_chunk)
                    temp_chunk = sentence
            if temp_chunk:
                # 仍然超长时，做硬切分兜底（避免embedding阶段失败）
                if len(temp_chunk) > max_length:
                    for i in range(0, len(temp_chunk), max_length):
                        chunks.append(temp_chunk[i : i + max_length])
                    current_chunk = ""
                else:
                    current_chunk = temp_chunk
        else:
            current_chunk = paragraph

    if current_chunk:
        chunks.append(current_chunk)
    return chunks

def _convert_doc_to_docx(doc_path: Path) -> Path:
    """将.doc格式转换为.docx格式（临时文件）"""
    try:
        import win32com.client as win32  # type: ignore
    except Exception as e:
        raise ValueError("处理.doc格式需要安装pywin32库") from e

    word = win32.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False
    doc = word.Documents.Open(str(doc_path.absolute()))

    temp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    temp_docx.close()

    try:
        doc.SaveAs(temp_docx.name, FileFormat=12)  # wdFormatXMLDocument = 12
        return Path(temp_docx.name)
    finally:
        try:
            doc.Close()
        except Exception as e:
            logger.warning(f"[rag] win32 word doc close failed: {e}", exc_info=True)
        try:
            word.Quit()
        except Exception as e:
            logger.warning(f"[rag] win32 word quit failed: {e}", exc_info=True)

def _read_word_file(path: Path) -> str:
    """读取.docx/.doc"""
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise ValueError("读取Word文档需要安装 python-docx") from e

    suffix = path.suffix.lower()
    temp_file: Optional[Path] = None
    try:
        if suffix == ".docx":
            doc = Document(path)
        elif suffix == ".doc":
            temp_file = _convert_doc_to_docx(path)
            doc = Document(temp_file)
        else:
            raise ValueError(f"不支持的Word文件类型: {suffix}")

        content_parts: List[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                content_parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    content_parts.append(" | ".join(row_text))

        return "\n".join(content_parts)
    finally:
        if temp_file and temp_file.exists():
            try:
                temp_file.unlink()
            except Exception as e:
                logger.warning(f"[rag] temp docx cleanup failed: {e}", exc_info=True)

def _read_file_content(file_path: str) -> str:
    """读取RAG入库支持的文件内容"""
    backend_path, _mapping = _apply_path_mapping(file_path)
    path = Path(backend_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在或后端无法访问: {file_path}")
    if not path.is_file():
        raise ValueError(f"不是文件: {file_path}")

    suffix = path.suffix.lower()
    if suffix in [".md", ".txt"]:
        return path.read_text(encoding="utf-8")
    if suffix in [".docx", ".doc"]:
        return _read_word_file(path)

    raise ValueError(f"不支持的文件类型: {suffix}")

def _list_supported_files(directory: str, recursive: bool = True, limit: int = 200) -> List[Dict[str, str]]:
    """List ingestible files under a directory for 'agent scan'."""
    backend_dir, _mapping = _apply_path_mapping(directory)
    p = Path(backend_dir)
    if not p.exists():
        raise FileNotFoundError(f"目录不存在: {directory}")
    if not p.is_dir():
        raise ValueError(f"不是目录: {directory}")

    supported = {".txt", ".md", ".docx", ".doc"}
    globber = p.rglob("*") if recursive else p.glob("*")
    out: List[Dict[str, str]] = []
    for fp in globber:
        if not fp.is_file():
            continue
        if fp.suffix.lower() not in supported:
            continue
        # Return backend-readable paths; UI may still show the original directory separately.
        out.append({"path": str(fp), "name": fp.name})
        if len(out) >= max(1, limit):
            break
    return out

def _upsert_file_to_vector_store(vector_store, file_path: str, import_method: str, document_name: str = "") -> Dict[str, Any]:
    """将文件解析、分块并写入向量库（存在则更新）"""
    from langchain_core.documents import Document  # 延迟导入

    # Canonicalize to a backend-readable path (shared-path mode), while keeping the original
    # path in metadata for display/debug.
    source_path, mapping = _apply_path_mapping(file_path)

    current_hash = _calculate_file_hash(source_path)
    prev_hash: Optional[str] = None
    file_exists = False

    # Prefer metadata-only probing via underlying collection to avoid loading chunk texts.
    try:
        collection = vector_store._store._collection
        probe = collection.get(where={"source": source_path}, include=["metadatas"], limit=1, offset=0)
        metas = probe.get("metadatas", []) or []
        if metas:
            file_exists = True
            prev_hash = (metas[0] or {}).get("file_hash") or (metas[0] or {}).get("hash")
    except Exception:
        # Fallback to store helper (may be heavier depending on implementation).
        try:
            file_exists = bool(vector_store.check_document_exists(source_path))
        except Exception:
            file_exists = False

    # Skip unchanged files (saves embedding time and avoids needless churn).
    if file_exists and current_hash and prev_hash and prev_hash == current_hash:
        return {
            "updated": False,
            "action": "skipped",
            "reason": "unchanged",
            "chunks_added": 0,
            "file_hash": current_hash,
            "source": source_path,
            "client_path": file_path,
        }

    content = _read_file_content(source_path)
    if not content.strip():
        raise ValueError("文件内容为空")

    if file_exists:
        try:
            vector_store.delete_by_filter({"source": source_path})
        except Exception as e:
            logger.warning(f"删除旧文档chunks失败 {source_path}: {e}")

    chunks = _split_into_chunks(content, max_length=1000)
    now = time.time()
    docs: List[Document] = []
    for i, chunk in enumerate(chunks):
        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "source": source_path,
                    "chunk_index": i,
                    "file_name": document_name or Path(source_path).name,
                    "import_method": import_method or "unknown",
                    "file_hash": current_hash,
                    "created_at": now,
                    "updated_at": now,
                    "client_path": file_path,
                    "path_mapping": mapping,
                },
            )
        )

    vector_store.add_documents(docs)
    return {
        "updated": True,
        "action": "updated" if file_exists else "added",
        "chunks_added": len(docs),
        "file_hash": current_hash,
        "prev_hash": prev_hash,
        "source": source_path,
        "client_path": file_path,
    }


def _sha256_text(text: str) -> str:
    try:
        h = hashlib.sha256()
        h.update((text or "").encode("utf-8", errors="ignore"))
        return h.hexdigest()
    except Exception:
        return ""


def _upsert_text_to_vector_store(
    vector_store,
    text: str,
    name: str,
    import_method: str = "upload",
    path_alias: Optional[str] = None,
) -> Dict[str, Any]:
    """Upload-to-RAG mode: embed plain text without requiring a backend-readable filesystem path."""
    from langchain_core.documents import Document  # 延迟导入

    content = (text or "").strip()
    if not content:
        raise ValueError("文本内容为空")

    sha = _sha256_text(content)
    if not sha:
        raise ValueError("无法计算文本哈希")

    source_ref = f"upload://{sha}"
    prev_hash: Optional[str] = None
    exists = False

    try:
        collection = vector_store._store._collection
        probe = collection.get(where={"source": source_ref}, include=["metadatas"], limit=1, offset=0)
        metas = probe.get("metadatas", []) or []
        if metas:
            exists = True
            prev_hash = (metas[0] or {}).get("file_hash") or (metas[0] or {}).get("hash")
    except Exception:
        exists = False

    if exists and prev_hash and prev_hash == sha:
        return {
            "updated": False,
            "action": "skipped",
            "reason": "unchanged",
            "chunks_added": 0,
            "file_hash": sha,
            "source": source_ref,
            "client_path": path_alias or name,
        }

    if exists:
        try:
            vector_store.delete_by_filter({"source": source_ref})
        except Exception as e:
            logger.warning(f"删除旧文档chunks失败 {source_ref}: {e}")

    chunks = _split_into_chunks(content, max_length=1000)
    now = time.time()
    docs: List[Document] = []
    for i, chunk in enumerate(chunks):
        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "source": source_ref,
                    "chunk_index": i,
                    "file_name": (name or "").strip() or "upload.txt",
                    "import_method": import_method or "upload",
                    "file_hash": sha,
                    "created_at": now,
                    "updated_at": now,
                    "client_path": (path_alias or name or "").strip() or source_ref,
                },
            )
        )

    vector_store.add_documents(docs)
    return {
        "updated": True,
        "action": "updated" if exists else "added",
        "chunks_added": len(docs),
        "file_hash": sha,
        "prev_hash": prev_hash,
        "source": source_ref,
        "client_path": path_alias or name,
    }


def _scan_directory_for_import(directory: str, recursive: bool = True, limit: int = 2000, import_method: str = "agent") -> List[BatchImportDocumentItem]:
    """Scan a directory for supported file types and return batch import items.

    Runs in the background task runner (can be slow for huge directories).
    """
    supported = {".txt", ".md", ".docx", ".doc"}
    backend_dir, _mapping = _apply_path_mapping(directory)
    directory_path = Path(backend_dir)
    if not directory_path.exists() or not directory_path.is_dir():
        raise ValueError(f"目录不存在或不是目录: {directory}")

    items: List[BatchImportDocumentItem] = []

    def _maybe_add(p: Path):
        if len(items) >= max(1, limit):
            return
        if p.is_file() and p.suffix.lower() in supported:
            items.append(
                BatchImportDocumentItem(path=str(p), name=p.name, importMethod=import_method)
            )

    if recursive:
        for p in directory_path.rglob("*"):
            _maybe_add(p)
            if len(items) >= max(1, limit):
                break
    else:
        for p in directory_path.iterdir():
            _maybe_add(p)
            if len(items) >= max(1, limit):
                break

    return items

def _compute_method_counts(metadatas: List[Dict[str, Any]]) -> Dict[str, int]:
    """统计每种导入方式的文档数量（按source去重）

    注意：只使用 metadata，避免加载所有 chunk 文本导致慢/超时。
    """
    buckets: Dict[str, set] = {}
    for meta in metadatas:
        meta = meta or {}
        source = meta.get("source", "unknown")
        method = meta.get("import_method", "unknown") or "unknown"
        buckets.setdefault(method, set()).add(source)
    return {k: len(v) for k, v in buckets.items()}

# API端点
@router.get("/statistics", response_model=RagStatisticsApiResponse)
async def get_statistics(
    projectId: Optional[str] = None,
    contextDocumentPath: Optional[str] = None,
    includeGlobal: bool = True,
    vector_store=Depends(get_vector_store),
):
    """获取RAG知识库统计信息。

    - 默认（无 projectId/contextDocumentPath）：返回全库统计
    - 传 projectId 或 contextDocumentPath：返回“当前项目(+全局)”范围内统计
    """
    try:
        idx = _get_project_index(vector_store)
        pid, plabel = _resolve_project_scope(
            vector_store=vector_store,
            project_id=projectId,
            context_document_path=contextDocumentPath,
        )

        scoped_sources: Optional[Set[str]] = None
        project_sources: List[str] = []
        global_sources: List[str] = []
        if pid:
            if pid == GLOBAL_PROJECT_ID:
                global_sources = idx.get_global_sources()
                scoped_sources = set(global_sources)
            else:
                project_sources = idx.get_project_sources(pid)
                global_sources = idx.get_global_sources() if includeGlobal else []
                scoped_sources = set(project_sources) | set(global_sources)

        total_documents = 0
        total_vectors = 0
        metadatas: List[Dict[str, Any]] = []

        if scoped_sources is None:
            stats = vector_store.get_stats()
            total_documents = int(stats.get("total_documents", 0) or 0)
            total_vectors = int(stats.get("total_chunks", 0) or 0)
            metadatas = vector_store.get_all_metadatas()
        else:
            collection = vector_store._store._collection
            where = {"source": {"$in": sorted(scoped_sources)}} if scoped_sources else {"source": "__none__"}
            res = collection.get(where=where, include=["metadatas"])
            metadatas = res.get("metadatas", []) or []
            total_vectors = len(metadatas)
            total_documents = len({(m or {}).get("source") for m in metadatas if (m or {}).get("source")})

        # 计算存储大小（估算，仍为全库磁盘占用）
        storage_path = vector_store.persist_path
        total_size = 0
        if storage_path.exists():
            for file_path in storage_path.rglob("*"):
                if file_path.is_file():
                    total_size += file_path.stat().st_size
        storage_size = _format_file_size(total_size)

        method_counts = _compute_method_counts(metadatas)

        def _method_info(key: str, name: str, icon: str) -> Dict[str, Any]:
            return {"count": method_counts.get(key, 0), "name": name, "icon": icon}

        scoped_doc_sources = {(m or {}).get("source") for m in metadatas if (m or {}).get("source")}
        global_doc_count = len(scoped_doc_sources & set(global_sources)) if pid else None
        project_doc_count = len(scoped_doc_sources & set(project_sources)) if pid else None

        return {
            "success": True,
            "data": {
                "totalDocuments": total_documents,
                "totalVectors": total_vectors,
                "storageSize": storage_size,
                "importMethods": {
                    "wps": _method_info("wps", "WPS同步导入", "📄"),
                    "agent": _method_info("agent", "Agent智能导入", "🤖"),
                    "api": _method_info("api", "API集成导入", "🔗"),
                    "command": _method_info("command", "命令行导入", "⌨️"),
                    "atref": _method_info("atref", "@引用导入", "📎"),
                },
                "scopeProjectId": pid,
                "scopeProjectLabel": plabel,
                "includeGlobal": bool(includeGlobal),
                "globalDocuments": global_doc_count,
                "projectDocuments": project_doc_count,
            },
        }
    except Exception as e:
        logger.error(f"获取统计信息失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取统计信息失败: {str(e)}")

@router.get("/documents/by-method", response_model=DocumentsByMethodApiResponse)
async def get_documents_by_method(
    projectId: Optional[str] = None,
    contextDocumentPath: Optional[str] = None,
    includeGlobal: bool = True,
    vector_store=Depends(get_vector_store),
):
    """按导入方式获取文档列表"""
    try:
        idx = _get_project_index(vector_store)
        pid, _plabel = _resolve_project_scope(
            vector_store=vector_store,
            project_id=projectId,
            context_document_path=contextDocumentPath,
        )

        metadatas: List[Dict[str, Any]] = []
        if pid:
            allowed = idx.get_allowed_sources(pid, include_global=bool(includeGlobal))
            if allowed:
                collection = vector_store._store._collection
                res = collection.get(where={"source": {"$in": allowed}}, include=["metadatas"])
                metadatas = res.get("metadatas", []) or []
            else:
                metadatas = []
        else:
            metadatas = vector_store.get_all_metadatas()

        method_names = {
            "wps": "WPS同步导入",
            "agent": "Agent智能导入",
            "api": "API集成导入",
            "command": "命令行导入",
            "atref": "@引用导入",
            "unknown": "未知导入方式",
        }

        # method -> source -> aggregated info
        grouped: Dict[str, Dict[str, Dict[str, Any]]] = {}
        for meta in metadatas:
            meta = meta or {}
            source = meta.get("source", "未知来源")
            if not source or source == "未知来源":
                continue
            method = meta.get("import_method", "unknown") or "unknown"
            grouped.setdefault(method, {})
            if source not in grouped[method]:
                file_stats = _get_file_stats(source)
                scope = idx.get_scope_for_source(source)
                projects_count = len(idx.get_projects_for_source(source))
                grouped[method][source] = {
                    "name": meta.get("file_name") or Path(source).name,
                    "path": source,
                    "displayPath": meta.get("client_path") or source,
                    "size": file_stats["size"],
                    "importTime": meta.get("created_at") or file_stats["modified"],
                    "vectors": 0,
                    "hash": meta.get("file_hash") or meta.get("hash"),
                    "scope": scope,
                    "projectsCount": projects_count,
                }
            grouped[method][source]["vectors"] += 1

        result: List[DocumentsByMethodResponse] = []
        for method, by_source in grouped.items():
            docs_list = [
                RagDocument(
                    name=info["name"],
                    path=info["path"],
                    displayPath=info.get("displayPath"),
                    size=info["size"],
                    importMethod=method,
                    importTime=str(info["importTime"]),
                    vectors=int(info["vectors"]),
                    hash=info.get("hash"),
                    scope=info.get("scope"),
                    projectsCount=info.get("projectsCount"),
                )
                for info in by_source.values()
            ]
            if docs_list:
                result.append(
                    DocumentsByMethodResponse(
                        method=method,
                        name=method_names.get(method, "未知导入方式"),
                        count=len(docs_list),
                        documents=docs_list,
                    )
                )

        order = {"wps": 0, "agent": 1, "api": 2, "command": 3, "atref": 4, "unknown": 99}
        result.sort(key=lambda x: order.get(x.method, 50))
        return {"success": True, "data": result}
    except Exception as e:
        logger.error(f"获取文档列表失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取文档列表失败: {str(e)}")

@router.delete("/documents")
async def delete_document(request: DeleteDocumentRequest, vector_store=Depends(get_vector_store)):
    """删除RAG文档"""
    try:
        client_path = request.documentPath
        document_path, _mapping = _apply_path_mapping(client_path)

        # 根据源文件路径删除文档
        deleted_count = vector_store.delete_by_filter({"source": document_path})

        if deleted_count > 0:
            # Also remove membership links so project/global views stay consistent.
            try:
                proj_idx = _get_project_index(vector_store)
                proj_idx.remove_source_everywhere(document_path)
                proj_idx.save()
            except Exception as e:
                logger.warning(f"[rag] delete: failed to update project index: {e}", exc_info=True)
            logger.info(f"成功删除文档 {document_path} (client={client_path}) 的 {deleted_count} 个chunks")
            return {
                "success": True,
                "message": f"成功删除文档",
                "deletedChunks": deleted_count
            }
        else:
            return {
                "success": False,
                "message": "未找到要删除的文档"
            }
    except Exception as e:
        logger.error(f"删除文档失败: {e}")
        raise HTTPException(status_code=500, detail=f"删除文档失败: {str(e)}")

@router.post("/documents/reimport")
async def reimport_document(request: ReimportDocumentRequest, vector_store=Depends(get_vector_store)):
    """重新导入文档"""
    try:
        document_path = request.documentPath
        document_name = request.documentName

        if str(document_path or "").startswith("upload://"):
            raise HTTPException(status_code=400, detail="upload:// 文档不支持重导入，请重新上传/再次同步入库")

        # 检查文件是否存在/可读（shared-path mode 会做映射）
        backend_path, _mapping = _apply_path_mapping(document_path)
        if not Path(backend_path).exists():
            raise HTTPException(status_code=404, detail=f"文件不存在或后端无法访问: {backend_path}")

        logger.info(f"重新导入文档: {document_name} ({document_path})")
        return {
            "success": True,
            "message": "文档重新导入成功",
            "data": _upsert_file_to_vector_store(
                vector_store=vector_store,
                file_path=document_path,
                import_method="wps",
                document_name=document_name,
            ),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"重新导入文档失败: {e}")
        raise HTTPException(status_code=500, detail=f"重新导入文档失败: {str(e)}")

@router.get("/documents")
async def get_all_documents(
    projectId: Optional[str] = None,
    contextDocumentPath: Optional[str] = None,
    includeGlobal: bool = True,
    vector_store=Depends(get_vector_store),
):
    """获取所有RAG文档"""
    try:
        # NOTE: 列表接口仅需 metadata，不要加载所有 chunk 文本（否则大库会非常慢）。
        idx = _get_project_index(vector_store)
        pid, _plabel = _resolve_project_scope(
            vector_store=vector_store,
            project_id=projectId,
            context_document_path=contextDocumentPath,
        )

        metadatas: List[Dict[str, Any]] = []
        if pid:
            allowed = idx.get_allowed_sources(pid, include_global=bool(includeGlobal))
            if allowed:
                collection = vector_store._store._collection
                res = collection.get(where={"source": {"$in": allowed}}, include=["metadatas"])
                metadatas = res.get("metadatas", []) or []
            else:
                metadatas = []
        else:
            metadatas = vector_store.get_all_metadatas()

        # source -> aggregated info
        by_source: Dict[str, Dict[str, Any]] = {}
        for meta in metadatas:
            meta = meta or {}
            source = meta.get("source", "")
            if not source:
                continue

            if source not in by_source:
                stats = _get_file_stats(source)
                scope = idx.get_scope_for_source(source)
                projects_count = len(idx.get_projects_for_source(source))
                by_source[source] = {
                    "name": meta.get("file_name") or Path(source).name,
                    "path": source,
                    "displayPath": meta.get("client_path") or source,
                    "size": stats["size"],
                    "importMethod": meta.get("import_method", "未知"),
                    "importTime": meta.get("created_at", "") or stats["modified"],
                    "vectors": 0,
                    "hash": meta.get("file_hash") or meta.get("hash"),
                    "scope": scope,
                    "projectsCount": projects_count,
                }

            by_source[source]["vectors"] += 1

        simplified_docs = list(by_source.values())

        return {
            "success": True,
            "data": simplified_docs
        }
    except Exception as e:
        logger.error(f"获取文档失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取文档失败: {str(e)}")

@router.get("/search")
async def search_documents(
    query: str,
    limit: int = 10,
    projectId: Optional[str] = None,
    contextDocumentPath: Optional[str] = None,
    includeGlobal: bool = True,
    vector_store=Depends(get_vector_store),
):
    """搜索RAG文档"""
    try:
        if not query.strip():
            return {"success": True, "data": []}

        idx = _get_project_index(vector_store)
        pid, _plabel = _resolve_project_scope(
            vector_store=vector_store,
            project_id=projectId,
            context_document_path=contextDocumentPath,
        )

        def _search_by_sources(sources: List[str], k: int) -> List[Any]:
            if not sources:
                return []
            # If the list is huge, Chroma may still handle it, but keep it bounded.
            bounded = sources[:5000]
            return vector_store.similarity_search(query, k=k, filter={"source": {"$in": bounded}})

        results: List[Any] = []
        project_sources: List[str] = []
        global_sources: List[str] = []
        if pid:
            if pid == GLOBAL_PROJECT_ID:
                global_sources = idx.get_global_sources()
                results = _search_by_sources(global_sources, limit)
            else:
                project_sources = idx.get_project_sources(pid)
                results = _search_by_sources(project_sources, limit)
                if includeGlobal and len(results) < limit:
                    global_sources = idx.get_global_sources()
                    results.extend(_search_by_sources(global_sources, limit))
        else:
            # No scoping information -> fallback to global search.
            results = vector_store.similarity_search(query, k=limit)

        # 转换为响应格式
        search_results = []
        seen = set()
        for doc, score in results:
            metadata = doc.metadata
            source = metadata.get("source", "")
            key = (source, metadata.get("chunk_index"), hash(doc.page_content))
            if key in seen:
                continue
            seen.add(key)
            # Enforce project-first ordering by truncating after de-dup.
            if len(search_results) >= int(limit or 10):
                break

            scope = idx.get_scope_for_source(source)
            search_results.append({
                "content": doc.page_content,
                "source": source,
                "score": score,
                "metadata": metadata,
                "scope": scope,
            })

        return {
            "success": True,
            "data": search_results
        }
    except Exception as e:
        logger.error(f"搜索文档失败: {e}")
        raise HTTPException(status_code=500, detail=f"搜索文档失败: {str(e)}")

@router.get("/documents/detail")
async def get_document_detail(documentPath: str, limit: int = 50, vector_store=Depends(get_vector_store)):
    """获取文档详情"""
    try:
        # 仅返回 chunk 预览，避免一次性拉取全量 chunk 文本导致卡顿/超时。
        collection = vector_store._store._collection
        client_path = documentPath
        source_path, _mapping = _apply_path_mapping(client_path)

        meta_res = collection.get(where={"source": source_path}, include=["metadatas"])
        total_chunks = len(meta_res.get("metadatas", []) or [])
        if total_chunks <= 0:
            raise HTTPException(status_code=404, detail=f"未找到文档: {source_path}")

        preview_limit = max(1, min(int(limit or 200), 200))
        preview_res = collection.get(
            where={"source": source_path},
            include=["metadatas", "documents"],
            limit=preview_limit,
            offset=0,
        )
        docs = []
        for i, (doc, meta) in enumerate(zip(preview_res.get("documents", []) or [], preview_res.get("metadatas", []) or [])):
            docs.append({
                "id": f"doc_{hash(doc)}_{i}",
                "content": doc,
                "content_preview": (doc[:200] + "...") if isinstance(doc, str) and len(doc) > 200 else doc,
                "metadata": meta,
                "created_at": (meta or {}).get("created_at", 0),
            })

        file_stats = _get_file_stats(source_path)
        first_meta = (meta_res.get("metadatas", []) or [{}])[0] or {}
        display_path = first_meta.get("client_path") or client_path or source_path

        return {
            "success": True,
            "data": {
                "name": first_meta.get("file_name") or Path(source_path).name,
                "path": source_path,
                "displayPath": display_path,
                "size": file_stats["size"],
                "modified": file_stats["modified"],
                "chunks": total_chunks,
                "chunksData": docs,
                "previewLimit": preview_limit,
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取文档详情失败: {e}")
        raise HTTPException(status_code=500, detail=f"获取文档详情失败: {str(e)}")

@router.post("/documents/batch-import")
async def batch_import_documents(request: BatchImportRequest, vector_store=Depends(get_vector_store)):
    """批量导入文档"""
    try:
        # Start an async ingestion task and return immediately (no fixed timeout on frontend).
        task_id = f"rag_{uuid.uuid4().hex[:12]}"
        state = RagTaskState(
            taskId=task_id,
            type="batch-import",
            status="pending",
            progress=0,
            currentStep="等待执行",
            createdAt=_now_ts(),
            updatedAt=_now_ts(),
        )
        _rag_tasks[task_id] = state
        _rag_task_queues[task_id] = asyncio.Queue()

        async def _runner():
            try:
                # Yield to event loop so the HTTP response can be sent before heavy work starts.
                await asyncio.sleep(0)
                _task_update(task_id, status="running", currentStep="开始入库", progress=1)
                proj_idx = _get_project_index(vector_store)
                scope = (request.scope or "").strip().lower()
                if scope not in ("global", "project"):
                    # Backward-compatible defaults:
                    # - directory scan imports (RAG中心) -> global
                    # - explicit documents list (legacy) -> global
                    scope = "global" if request.directory else "global"

                docs = request.documents or []
                if (not docs) and request.directory:
                    _task_update(task_id, currentStep="扫描目录中...", progress=3)
                    # Scan in a worker thread to avoid blocking the event loop (prevents request timeouts).
                    docs = await asyncio.to_thread(
                        _scan_directory_for_import,
                        directory=request.directory,
                        recursive=bool(request.recursive),
                        limit=int(request.limit or 2000),
                        import_method=request.importMethod or "agent",
                    )
                    if not docs:
                        raise ValueError("未扫描到可导入的文档（当前仅支持 txt/md/doc/docx）")
                    _task_update(task_id, currentStep=f"扫描完成，发现 {len(docs)} 个文档，开始入库", progress=8)

                if not docs:
                    raise ValueError("documents 为空，且未提供 directory")
                total = max(len(docs), 1)
                success_count = 0
                failed_count = 0
                results: List[Dict[str, Any]] = []

                for i, doc_info in enumerate(docs):
                    # Cancellation check
                    if task_id in _rag_async_tasks and _rag_async_tasks[task_id].cancelled():
                        raise asyncio.CancelledError()

                    _task_update(
                        task_id,
                        currentStep=f"入库中: {doc_info.name} ({i+1}/{len(docs)})",
                        progress=min(99, int((i / total) * 90) + 5),
                    )
                    try:
                        # Ingestion (read + embed + write) is CPU/IO heavy; keep it off the event loop.
                        if doc_info.text and str(doc_info.text).strip():
                            ingest_result = await asyncio.to_thread(
                                _upsert_text_to_vector_store,
                                vector_store,
                                doc_info.text,
                                doc_info.name,
                                doc_info.importMethod,
                                doc_info.pathAlias or doc_info.path or None,
                            )
                        else:
                            ingest_result = await asyncio.to_thread(
                                _upsert_file_to_vector_store,
                                vector_store,
                                doc_info.path,
                                doc_info.importMethod,
                                doc_info.name,
                            )
                        success_count += 1
                        results.append({"success": True, "path": doc_info.path, "name": doc_info.name, "data": ingest_result})

                        # Update project/global membership (does not duplicate vectors).
                        try:
                            src = str((ingest_result or {}).get("source") or doc_info.path or "")
                            if src:
                                if scope == "global":
                                    proj_idx.add_global(src)
                                else:
                                    # Derive project id from the file path by default.
                                    ctx = derive_project_context(src)
                                    if ctx:
                                        proj_idx.add_project(src, ctx)
                        except Exception as e:
                            logger.warning(f"[rag] update project index membership failed: {e}", exc_info=True)
                    except Exception as e:
                        failed_count += 1
                        results.append({"success": False, "path": doc_info.path, "name": doc_info.name, "error": str(e)})

                # Persist membership updates once per task.
                try:
                    proj_idx.save()
                except Exception as e:
                    logger.warning(f"[rag] project index save failed: {e}", exc_info=True)

                _task_update(
                    task_id,
                    status="completed",
                    progress=100,
                    currentStep="入库完成",
                    result={
                        "successCount": success_count,
                        "failedCount": failed_count,
                        "totalCount": len(docs),
                        "results": results,
                    },
                )
            except asyncio.CancelledError:
                _task_update(task_id, status="cancelled", currentStep="已取消", progress=0)
            except Exception as e:
                _task_update(task_id, status="failed", currentStep="入库失败", error=str(e))

        _rag_async_tasks[task_id] = asyncio.create_task(_runner())

        return {
            "success": True,
            "message": "已创建入库任务",
            "data": {"taskId": task_id, "task": state.model_dump()},
        }
    except Exception as e:
        logger.error(f"批量导入失败: {e}")
        raise HTTPException(status_code=500, detail=f"批量导入失败: {str(e)}")


async def _read_upload_file_content(upload: UploadFile) -> Tuple[str, str]:
    """Return (name, extracted_text) for an uploaded file."""
    name = (upload.filename or "upload.txt").strip() or "upload.txt"
    raw = await upload.read()
    suffix = Path(name).suffix.lower()

    if suffix in (".txt", ".md", ""):
        try:
            return name, raw.decode("utf-8", errors="ignore")
        except Exception:
            return name, raw.decode(errors="ignore")

    if suffix == ".docx":
        # Parse .docx by writing to a temp file for python-docx.
        tmp = None
        try:
            fd, p = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            tmp = Path(p)
            tmp.write_bytes(raw)
            return name, _read_word_file(tmp)
        finally:
            if tmp and tmp.exists():
                try:
                    tmp.unlink()
                except Exception as e:
                    logger.warning(f"[rag] upload temp cleanup failed: {e}", exc_info=True)

    if suffix == ".doc":
        # .doc conversion requires Windows + Word automation; keep this explicit.
        raise ValueError(".doc 不支持上传解析（需要 Windows + Word 组件），请转换为 .docx")

    raise ValueError(f"不支持的上传文件类型: {suffix}")


@router.post("/documents/upload-text")
async def upload_text_documents(request: UploadTextRequest, vector_store=Depends(get_vector_store)):
    """Upload-to-RAG: client sends text, backend embeds without filesystem access."""
    try:
        task_id = f"rag_{uuid.uuid4().hex[:12]}"
        state = RagTaskState(
            taskId=task_id,
            type="upload-text",
            status="pending",
            progress=0,
            currentStep="等待执行",
            createdAt=_now_ts(),
            updatedAt=_now_ts(),
        )
        _rag_tasks[task_id] = state
        _rag_task_queues[task_id] = asyncio.Queue()

        async def _runner():
            try:
                await asyncio.sleep(0)
                _task_update(task_id, status="running", currentStep="开始入库", progress=1)
                proj_idx = _get_project_index(vector_store)

                items = request.items or []
                if not items:
                    raise ValueError("items 为空")

                total = max(len(items), 1)
                success_count = 0
                failed_count = 0
                results: List[Dict[str, Any]] = []

                for i, item in enumerate(items):
                    if task_id in _rag_async_tasks and _rag_async_tasks[task_id].cancelled():
                        raise asyncio.CancelledError()

                    _task_update(
                        task_id,
                        currentStep=f"入库中: {item.name} ({i+1}/{len(items)})",
                        progress=min(99, int((i / total) * 90) + 5),
                    )
                    try:
                        ingest_result = await asyncio.to_thread(
                            _upsert_text_to_vector_store,
                            vector_store,
                            item.text,
                            item.name,
                            item.importMethod,
                            item.pathAlias or item.name,
                        )
                        success_count += 1
                        results.append({"success": True, "name": item.name, "data": ingest_result})

                        # Membership: global vs project (project derived from contextDocumentPath).
                        try:
                            src = str((ingest_result or {}).get("source") or "")
                            if not src:
                                continue
                            scope = (item.scope or "global").strip().lower()
                            if scope == "project":
                                ctx = derive_project_context(item.contextDocumentPath or "")
                                if ctx:
                                    proj_idx.add_project(src, ctx)
                                else:
                                    proj_idx.add_global(src)
                            else:
                                proj_idx.add_global(src)
                        except Exception as e:
                            logger.warning(f"[rag] update project index membership failed: {e}", exc_info=True)
                    except Exception as e:
                        failed_count += 1
                        results.append({"success": False, "name": item.name, "error": str(e)})

                try:
                    proj_idx.save()
                except Exception as e:
                    logger.warning(f"[rag] project index save failed: {e}", exc_info=True)

                _task_update(
                    task_id,
                    status="completed",
                    progress=100,
                    currentStep="入库完成",
                    result={
                        "successCount": success_count,
                        "failedCount": failed_count,
                        "totalCount": len(items),
                        "results": results,
                    },
                )
            except asyncio.CancelledError:
                _task_update(task_id, status="cancelled", currentStep="已取消", progress=0)
            except Exception as e:
                _task_update(task_id, status="failed", currentStep="入库失败", error=str(e))

        _rag_async_tasks[task_id] = asyncio.create_task(_runner())
        return {"success": True, "message": "已创建入库任务", "data": {"taskId": task_id, "task": state.model_dump()}}
    except Exception as e:
        logger.error(f"upload-text 入库失败: {e}")
        raise HTTPException(status_code=500, detail=f"upload-text 入库失败: {str(e)}")


@router.post("/documents/upload-files")
async def upload_files_documents(
    files: List[UploadFile] = File(...),
    scope: str = Form("global"),
    contextDocumentPath: Optional[str] = Form(None),
    importMethod: str = Form("upload"),
    vector_store=Depends(get_vector_store),
):
    """Upload-to-RAG: multipart file upload -> text extract -> embed."""
    try:
        task_id = f"rag_{uuid.uuid4().hex[:12]}"
        state = RagTaskState(
            taskId=task_id,
            type="upload-files",
            status="pending",
            progress=0,
            currentStep="等待执行",
            createdAt=_now_ts(),
            updatedAt=_now_ts(),
        )
        _rag_tasks[task_id] = state
        _rag_task_queues[task_id] = asyncio.Queue()

        async def _runner():
            try:
                await asyncio.sleep(0)
                _task_update(task_id, status="running", currentStep="开始入库", progress=1)
                proj_idx = _get_project_index(vector_store)

                total = max(len(files or []), 1)
                success_count = 0
                failed_count = 0
                results: List[Dict[str, Any]] = []

                for i, f in enumerate(files or []):
                    if task_id in _rag_async_tasks and _rag_async_tasks[task_id].cancelled():
                        raise asyncio.CancelledError()

                    _task_update(
                        task_id,
                        currentStep=f"解析/入库中: {(f.filename or 'upload')} ({i+1}/{len(files)})",
                        progress=min(99, int((i / total) * 90) + 5),
                    )
                    try:
                        name, text = await _read_upload_file_content(f)
                        ingest_result = await asyncio.to_thread(
                            _upsert_text_to_vector_store,
                            vector_store,
                            text,
                            name,
                            importMethod,
                            name,
                        )
                        success_count += 1
                        results.append({"success": True, "name": name, "data": ingest_result})

                        try:
                            src = str((ingest_result or {}).get("source") or "")
                            if not src:
                                continue
                            if (scope or "global").strip().lower() == "project":
                                ctx = derive_project_context(contextDocumentPath or "")
                                if ctx:
                                    proj_idx.add_project(src, ctx)
                                else:
                                    proj_idx.add_global(src)
                            else:
                                proj_idx.add_global(src)
                        except Exception as e:
                            logger.warning(f"[rag] update project index membership failed: {e}", exc_info=True)
                    except Exception as e:
                        failed_count += 1
                        results.append({"success": False, "name": (f.filename or "upload"), "error": str(e)})

                try:
                    proj_idx.save()
                except Exception as e:
                    logger.warning(f"[rag] project index save failed: {e}", exc_info=True)

                _task_update(
                    task_id,
                    status="completed",
                    progress=100,
                    currentStep="入库完成",
                    result={
                        "successCount": success_count,
                        "failedCount": failed_count,
                        "totalCount": len(files or []),
                        "results": results,
                    },
                )
            except asyncio.CancelledError:
                _task_update(task_id, status="cancelled", currentStep="已取消", progress=0)
            except Exception as e:
                _task_update(task_id, status="failed", currentStep="入库失败", error=str(e))

        _rag_async_tasks[task_id] = asyncio.create_task(_runner())
        return {"success": True, "message": "已创建入库任务", "data": {"taskId": task_id, "task": state.model_dump()}}
    except Exception as e:
        logger.error(f"upload-files 入库失败: {e}")
        raise HTTPException(status_code=500, detail=f"upload-files 入库失败: {str(e)}")

@router.get("/tasks")
async def get_import_tasks():
    """获取导入任务状态"""
    return {
        "success": True,
        "data": [t.model_dump() for t in _rag_tasks.values()]
    }

@router.get("/tasks/{taskId}")
async def get_import_task_status(taskId: str):
    task = _rag_tasks.get(taskId)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")
    return {"success": True, "data": task.model_dump()}

@router.get("/tasks/{taskId}/stream")
async def stream_import_task(taskId: str):
    """SSE stream for task updates (event-driven, no polling)."""
    task = _rag_tasks.get(taskId)
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    q = _rag_task_queues.get(taskId)
    if not q:
        _rag_task_queues[taskId] = asyncio.Queue()
        q = _rag_task_queues[taskId]

    async def event_gen():
        # send initial snapshot
        yield f"event: task\ndata: {task.model_dump_json()}\n\n"
        while True:
            try:
                msg = await asyncio.wait_for(q.get(), timeout=15.0)
                ev_type = msg.get("type", "task")
                data = msg.get("data", {})
                # Note: data may already be JSON-serializable.
                import json
                yield f"event: {ev_type}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

                # Stop streaming on terminal state.
                state = _rag_tasks.get(taskId)
                if state and state.status in ("completed", "failed", "cancelled"):
                    break
            except asyncio.TimeoutError:
                # keepalive ping
                yield "event: ping\ndata: {}\n\n"

    return StreamingResponse(event_gen(), media_type="text/event-stream")

@router.post("/tasks/{taskId}/cancel")
async def cancel_import_task(taskId: str):
    """取消导入任务"""
    t = _rag_async_tasks.get(taskId)
    if t and not t.done():
        t.cancel()
    _task_update(taskId, status="cancelled", currentStep="已取消", progress=0)
    return {
        "success": True,
        "message": "任务已取消"
    }

@router.delete("/knowledge-base")
async def clear_knowledge_base(vector_store=Depends(get_vector_store)):
    """清空RAG知识库"""
    try:
        vector_store.clear()
        logger.info("已清空RAG知识库")
        return {
            "success": True,
            "message": "知识库已清空"
        }
    except Exception as e:
        logger.error(f"清空知识库失败: {e}")
        raise HTTPException(status_code=500, detail=f"清空知识库失败: {str(e)}")

@router.get("/config")
async def get_knowledge_base_config():
    """获取知识库配置信息"""
    # TODO: 实现配置管理
    return {
        "success": True,
        "data": {
            "embeddingModel": settings.embedding_model,
            "chunkSize": 1000,
            "chunkOverlap": 200,
            "maxTokens": 4000
        }
    }

@router.put("/config")
async def update_knowledge_base_config(config: Dict[str, Any]):
    """更新知识库配置"""
    # TODO: 实现配置更新
    return {
        "success": True,
        "message": "配置已更新"
    }

@router.get("/history")
async def get_import_history(limit: int = 50):
    """获取导入历史"""
    # TODO: 实现历史记录
    return {
        "success": True,
        "data": []
    }

@router.get("/export")
async def export_documents(
    format: str = "json",
    projectId: Optional[str] = None,
    contextDocumentPath: Optional[str] = None,
    includeGlobal: bool = True,
    includeIndex: bool = True,
    batchSize: int = 200,
    vector_store=Depends(get_vector_store),
):
    """导出RAG文档（用于迁移/同步）。

    - format=json: 返回一个 JSON payload（适合小库/调试）
    - format=jsonl: 逐行输出 JSONL（适合大库）

    支持按项目范围导出：projectId 或 contextDocumentPath（推导项目）+ includeGlobal。
    """

    def _iter_collection_docs(collection, where: Optional[Dict[str, Any]], batch_size: int):
        offset = 0
        bs = max(1, min(int(batch_size or 200), 2000))
        while True:
            res = collection.get(
                where=where,
                include=["documents", "metadatas"],
                limit=bs,
                offset=offset,
            )
            docs = res.get("documents") or []
            metas = res.get("metadatas") or []
            if not docs:
                break
            for doc, meta in zip(docs, metas):
                yield doc, (meta or {})
            offset += len(docs)
            if len(docs) < bs:
                break

    fmt = (format or "json").strip().lower()
    if fmt not in {"json", "jsonl"}:
        raise HTTPException(status_code=400, detail=f"unsupported format: {format}")

    idx = _get_project_index(vector_store)
    pid, plabel = _resolve_project_scope(
        vector_store=vector_store,
        project_id=projectId,
        context_document_path=contextDocumentPath,
    )
    allowed_sources: Optional[Set[str]] = None
    if pid:
        allowed_sources = set(idx.get_allowed_sources(pid, include_global=includeGlobal))

    try:
        collection = vector_store._store._collection
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"vector store not ready: {e}") from e

    where: Optional[Dict[str, Any]] = None
    if allowed_sources is not None:
        where = {"source": {"$in": sorted(allowed_sources)}} if allowed_sources else {"source": "__none__"}

    exported_at = time.time()

    project_index_payload: Optional[Dict[str, Any]] = None
    if includeIndex:
        try:
            # Ensure we export the latest snapshot (best-effort).
            idx.save()
            if idx.index_path.exists():
                project_index_payload = json.loads(idx.index_path.read_text(encoding="utf-8") or "{}")
            else:
                project_index_payload = {"version": 1, "global_sources": [], "projects": {}}
        except Exception:
            project_index_payload = None

    scope_payload = {
        "project_id": pid,
        "project_label": plabel,
        "include_global": bool(includeGlobal),
    }

    if fmt == "jsonl":
        def _stream():
            manifest = {
                "type": "manifest",
                "schema_version": "ah32.rag.export.v1",
                "exported_at": exported_at,
                "scope": scope_payload,
                "project_index": project_index_payload,
            }
            yield (json.dumps(manifest, ensure_ascii=False) + "\n").encode("utf-8")
            for content, meta in _iter_collection_docs(collection, where, batchSize):
                item = {
                    "type": "chunk",
                    "schema_version": "ah32.rag.export.v1",
                    "content": content,
                    "metadata": meta,
                }
                yield (json.dumps(item, ensure_ascii=False) + "\n").encode("utf-8")

        return StreamingResponse(_stream(), media_type="application/jsonl; charset=utf-8")

    # json mode (grouped by source)
    grouped: Dict[str, Dict[str, Any]] = {}
    total_chunks = 0
    for content, meta in _iter_collection_docs(collection, where, batchSize):
        total_chunks += 1
        src = str(meta.get("source") or "unknown")
        bucket = grouped.get(src)
        if not bucket:
            bucket = {
                "source": src,
                "file_hash": str(meta.get("file_hash") or meta.get("hash") or ""),
                "file_name": str(meta.get("file_name") or ""),
                "import_method": str(meta.get("import_method") or ""),
                "client_path": meta.get("client_path"),
                "path_mapping": meta.get("path_mapping"),
                "created_at": float(meta.get("created_at") or 0.0),
                "updated_at": float(meta.get("updated_at") or 0.0),
                "chunks": [],
            }
            grouped[src] = bucket
        bucket["chunks"].append(
            {
                "chunk_index": int(meta.get("chunk_index") or 0),
                "content": content,
            }
        )

    documents = list(grouped.values())
    for doc in documents:
        try:
            doc["chunks"] = sorted(doc.get("chunks") or [], key=lambda x: int(x.get("chunk_index") or 0))
        except Exception as e:
            logger.warning(f"[rag] sort export chunks failed: {e}", exc_info=True)

    return {
        "success": True,
        "data": {
            "schema_version": "ah32.rag.export.v1",
            "exported_at": exported_at,
            "scope": scope_payload,
            "project_index": project_index_payload,
            "documents": documents,
            "stats": {
                "sources": len(documents),
                "chunks": total_chunks,
            },
        },
    }


class RagExportChunk(BaseModel):
    chunk_index: int
    content: str


class RagExportDocument(BaseModel):
    source: str
    file_hash: str = ""
    file_name: str = ""
    import_method: str = ""
    client_path: Optional[str] = None
    path_mapping: Optional[Dict[str, Any]] = None
    created_at: float = 0.0
    updated_at: float = 0.0
    chunks: List[RagExportChunk]


class RagExportPayload(BaseModel):
    schema_version: str
    exported_at: float
    scope: Optional[Dict[str, Any]] = None
    project_index: Optional[Dict[str, Any]] = None
    documents: List[RagExportDocument]


@router.post("/import")
async def import_documents(
    request: Request,
    incremental: bool = True,
    mergeIndex: bool = True,
    vector_store=Depends(get_vector_store),
):
    """导入RAG文档（用于迁移/同步），支持增量（按 source+file_hash 去重）。"""
    try:
        from langchain_core.documents import Document
    except ImportError:  # pragma: no cover
        from langchain.documents import Document

    async def _load_payload() -> RagExportPayload:
        ct = (request.headers.get("content-type") or "").lower()
        if "multipart/form-data" in ct:
            form = await request.form()
            upload = form.get("file")
            if not upload:
                raise HTTPException(status_code=400, detail="missing file in multipart form")
            raw = await upload.read()
            text = raw.decode("utf-8", errors="ignore")
            try:
                data = json.loads(text)
            except Exception:
                data = None
            if isinstance(data, dict):
                return RagExportPayload.model_validate(data)
            # jsonl fallback
            docs: Dict[str, Dict[str, Any]] = {}
            manifest: Dict[str, Any] = {}
            for line in text.splitlines():
                line = line.strip()
                if not line:
                    continue
                item = json.loads(line)
                if item.get("type") == "manifest":
                    manifest = item
                    continue
                if item.get("type") != "chunk":
                    continue
                meta = item.get("metadata") or {}
                content = item.get("content") or ""
                src = str(meta.get("source") or "unknown")
                bucket = docs.get(src)
                if not bucket:
                    bucket = {
                        "source": src,
                        "file_hash": str(meta.get("file_hash") or meta.get("hash") or ""),
                        "file_name": str(meta.get("file_name") or ""),
                        "import_method": str(meta.get("import_method") or ""),
                        "client_path": meta.get("client_path"),
                        "path_mapping": meta.get("path_mapping"),
                        "created_at": float(meta.get("created_at") or 0.0),
                        "updated_at": float(meta.get("updated_at") or 0.0),
                        "chunks": [],
                    }
                    docs[src] = bucket
                bucket["chunks"].append(
                    {
                        "chunk_index": int(meta.get("chunk_index") or 0),
                        "content": content,
                    }
                )
            payload = {
                "schema_version": manifest.get("schema_version") or "ah32.rag.export.v1",
                "exported_at": float(manifest.get("exported_at") or time.time()),
                "scope": manifest.get("scope"),
                "project_index": manifest.get("project_index"),
                "documents": list(docs.values()),
            }
            return RagExportPayload.model_validate(payload)

        data = await request.json()
        if isinstance(data, dict) and "data" in data and isinstance(data.get("data"), dict):
            # Allow importing directly from the export API response.
            data = data.get("data")
        return RagExportPayload.model_validate(data)

    payload = await _load_payload()
    if str(payload.schema_version or "").strip() != "ah32.rag.export.v1":
        raise HTTPException(status_code=400, detail=f"unsupported schema_version: {payload.schema_version}")

    idx = _get_project_index(vector_store)
    if mergeIndex and payload.project_index:
        try:
            global_sources = payload.project_index.get("global_sources") or []
            if isinstance(global_sources, list):
                for s in global_sources:
                    if s:
                        idx.add_global(str(s))
            projects = payload.project_index.get("projects") or {}
            if isinstance(projects, dict):
                for pid, meta in projects.items():
                    if not pid or not isinstance(meta, dict):
                        continue
                    sources = meta.get("sources") or []
                    if not isinstance(sources, list):
                        continue
                    ctx = ProjectContext(
                        project_id=str(pid),
                        root=str(meta.get("root") or ""),
                        label=str(meta.get("label") or ""),
                    )
                    for s in sources:
                        if s:
                            idx.add_project(str(s), ctx)
            idx.save()
        except Exception as e:
            logger.warning(f"[rag] import: project index rebuild/save failed: {e}", exc_info=True)

    try:
        collection = vector_store._store._collection
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"vector store not ready: {e}") from e

    def _probe_hash(source: str) -> Tuple[bool, str]:
        try:
            res = collection.get(where={"source": source}, include=["metadatas"], limit=1, offset=0)
            metas = res.get("metadatas") or []
            if not metas:
                return False, ""
            meta0 = metas[0] or {}
            return True, str(meta0.get("file_hash") or meta0.get("hash") or "")
        except Exception:
            return False, ""

    def _upsert_export_doc(doc: RagExportDocument) -> Dict[str, Any]:
        src = str(doc.source or "").strip()
        if not src:
            return {"source": src, "action": "skipped", "reason": "missing_source", "chunks_added": 0}
        target_hash = str(doc.file_hash or "").strip()
        exists, prev_hash = _probe_hash(src)
        if incremental and exists and target_hash and prev_hash and prev_hash == target_hash:
            return {"source": src, "action": "skipped", "reason": "unchanged", "chunks_added": 0, "file_hash": target_hash}

        if exists:
            try:
                vector_store.delete_by_filter({"source": src})
            except Exception as e:
                logger.warning(f"[rag] import: failed to delete existing source before upsert: {e}", exc_info=True)

        now = time.time()
        created_at = float(doc.created_at or 0.0) or now
        updated_at = float(doc.updated_at or 0.0) or now
        docs: List[Document] = []
        for chunk in doc.chunks:
            docs.append(
                Document(
                    page_content=chunk.content,
                    metadata={
                        "source": src,
                        "chunk_index": int(chunk.chunk_index),
                        "file_name": str(doc.file_name or ""),
                        "import_method": str(doc.import_method or "sync"),
                        "file_hash": target_hash,
                        "created_at": created_at,
                        "updated_at": updated_at,
                        "client_path": doc.client_path,
                        "path_mapping": doc.path_mapping,
                        "synced_at": now,
                    },
                )
            )

        vector_store.add_documents(docs)
        return {
            "source": src,
            "action": "updated" if exists else "added",
            "chunks_added": len(docs),
            "file_hash": target_hash,
            "prev_hash": prev_hash,
        }

    results: List[Dict[str, Any]] = []
    stats = {"added": 0, "updated": 0, "skipped": 0, "chunks_added": 0}
    for doc in payload.documents or []:
        r = await asyncio.to_thread(_upsert_export_doc, doc)
        results.append(r)
        action = str(r.get("action") or "")
        if action in stats:
            stats[action] += 1
        stats["chunks_added"] += int(r.get("chunks_added") or 0)

    return {"success": True, "data": {"results": results, "stats": stats}}

@router.post("/validate-path")
async def validate_document_path(request: ValidatePathRequest):
    """验证文档路径"""
    try:
        path = request.path
        backend_path, mapping = _apply_path_mapping(path)
        file_path = Path(backend_path)
        if not file_path.exists():
            # Provide actionable hints for remote deployments.
            if mapping:
                reason = "映射后的路径不存在或后端无权限访问（请检查挂载/权限）"
            else:
                reason = "后端无法访问该路径（远程部署请配置 AH32_PATH_MAPPINGS 或使用上传入库）"
            return {
                "success": True,
                "data": {
                    "valid": False,
                    "error": reason,
                    "mappedPath": backend_path,
                    "mapping": mapping,
                }
            }

        if not file_path.is_file():
            return {
                "success": True,
                "data": {
                    "valid": False,
                    "error": "不是文件",
                    "mappedPath": backend_path,
                    "mapping": mapping,
                }
            }

        return {
            "success": True,
            "data": {
                "valid": True,
                "mappedPath": backend_path,
                "mapping": mapping,
            }
        }
    except Exception as e:
        return {
            "success": True,
            "data": {
                "valid": False,
                "error": str(e)
            }
        }

@router.get("/supported-formats")
async def get_supported_formats():
    """获取支持的文档格式"""
    return {
        "success": True,
        "data": [
            ".txt", ".md", ".docx", ".pdf", ".pptx",
            ".xlsx", ".csv", ".html", ".htm"
        ]
    }


@router.post("/scan")
async def scan_directory(request: ScanDirectoryRequest):
    """扫描目录，返回可入库文件列表"""
    try:
        docs = _list_supported_files(request.directory, recursive=request.recursive, limit=request.limit)
        return {
            "success": True,
            "data": {
                "documentCount": len(docs),
                "documents": docs,
            },
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
